import os
from pdf2docx import Converter
from docx import Document
from fpdf import FPDF
import pandas as pd
from PIL import Image

def convert_file(file_path, target_format):
    base, ext = os.path.splitext(file_path)
    ext = ext.lower().strip(".")
    converted_path = f"{base}.{target_format}"

    # PDF to DOCX
    if ext == "pdf" and target_format == "docx":
        cv = Converter(file_path)
        cv.convert(converted_path)
        cv.close()

    # DOCX to PDF
    elif ext == "docx" and target_format == "pdf":
        doc = Document(file_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        for para in doc.paragraphs:
            if pdf.get_y() > 270:  # Check if near bottom of page
                pdf.add_page()  # Add a new page if needed
            pdf.multi_cell(0, 10, txt=para.text)

        pdf.output(converted_path)

    # TXT to PDF
    elif ext == "txt" and target_format == "pdf":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        with open(file_path, "r", encoding="utf-8") as file:
            for line in file:
                if pdf.get_y() > 270:
                    pdf.add_page()
                pdf.multi_cell(0, 10, txt=line.strip())

        pdf.output(converted_path)

    # PNG to PDF
    elif ext == "png" and target_format == "pdf":
        img = Image.open(file_path)
        img.convert("RGB").save(converted_path)

    # XLSX to PDF
    elif ext == "xlsx" and target_format == "pdf":
        df = pd.read_excel(file_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=8)

        for row in df.itertuples(index=False):
            row_text = " | ".join(str(cell) for cell in row)
            if pdf.get_y() > 270:
                pdf.add_page()
            pdf.multi_cell(0, 10, txt=row_text)

        pdf.output(converted_path)

    # XLSX to DOCX
    elif ext == "xlsx" and target_format == "docx":
        df = pd.read_excel(file_path)
        doc = Document()
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = str(column_name)

        for row in df.itertuples(index=False):
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

        doc.save(converted_path)

    # CSV to PDF
    elif ext == "csv" and target_format == "pdf":
        df = pd.read_csv(file_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=8)

        for row in df.itertuples(index=False):
            row_text = " | ".join(str(cell) for cell in row)
            if pdf.get_y() > 270:
                pdf.add_page()
            pdf.multi_cell(0, 10, txt=row_text)

        pdf.output(converted_path)

    # CSV to DOCX
    elif ext == "csv" and target_format == "docx":
        df = pd.read_csv(file_path)
        doc = Document()
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = str(column_name)

        for row in df.itertuples(index=False):
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

        doc.save(converted_path)

    return converted_path
